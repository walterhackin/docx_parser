# -*- coding: utf-8 -*-

# -- Sheet --

import docx
import pandas as pd
from docx import Document
from tqdm import tqdm

DOCUMENT_PATH = ''

df = Document(DOCUMENT_PATH)

parsed_parts = []
for i in tqdm(range(25,len(df.paragraphs))):
    if df.paragraphs[i].text == "" or df.paragraphs[i].text == u'\xa0' or df.paragraphs[i].text == "  ":
        continue
    elif df.paragraphs[i].text.startswith(('(', ' (', ', ', ')')):
        parsed_parts[-1][1] = (df.paragraphs[i].text)
    elif df.paragraphs[i].text == "АЛФАВИТНЫЙ УКАЗАТЕЛЬ":
        break
    else:
        parsed_parts.append([df.paragraphs[i].text, ''])

def to_single(data):
    new_data = []
    new_data.append(data[0])
    new_data.append(data[1])
    for i in data[2:]:
        i = i.replace('\n', ' ~ ')
        if i not in new_data or i in ['', ' ', '  ', u'\xa0']:
            new_data.append(i)
    return new_data

personal_data = []

for tb in tqdm(df.tables[3:]):
    table_data = []
    for row in tb.rows:
        data = to_single([c.text for c in row.cells])
        if len(data) == 4:
            table_data.append([data[0], ' ', data[1], data[2], data[3]])
        elif len(data) == 3:
            table_data.append([data[0], ' ', data[1], data[2]])
        else:
            table_data.append([data[0], data[1], data[2], data[3], data[4]])
    personal_data.append(table_data)

def fill_dataframe(reference, personal_data, parsed_parts):
    for table_idx in range(len(personal_data)):
        for row in personal_data[table_idx]:
            reference['ФИО'].append(row[0])
            reference['Должность'].append(row[1])
            reference['Внутренний телефон-1'].append(row[2])
            reference['Внутренний телефон-2'].append('')
            reference['Внутренний телефон-3'].append('')
            reference['Телефон-1'].append(row[3])
            reference['Кабинет'].append(row[4])
            reference['Телефон-2'].append('')
            reference['Телефон-3'].append('')
            reference['Почта'].append('')
            department_data = parsed_parts[table_idx][0].split('->')
            departments = ['' for k in range(4)]
            departments[:len(department_data)] = parsed_parts[table_idx][0].split('->')
            reference['Адрес'].append(parsed_parts[table_idx][1] if len(parsed_parts[table_idx][1]) != 0 else ' ')
            department_names = list(reference.keys())[11:]
            for j in range(len(departments)):
                reference[department_names[j]].append(departments[j])

reference = {
                'ФИО' : [], 'Должность' : [], 'Внутренний телефон-1' : [], 'Внутренний телефон-2' : [],
                'Внутренний телефон-3' : [],
                'Телефон-1' : [], 'Телефон-2' : [], 'Телефон-3' : [], 'Почта': [], 'Кабинет' : [], 'Адрес' : [],
                'Подразделение-1': [], 'Подразделение-2': [], 'Подразделение-3': [] ,
                'Подразделение-4': []
}

fill_dataframe(reference, personal_data, parsed_parts)

reference_df = pd.DataFrame.from_dict(reference)

reference_df

parsed_numbers = ['Телефон-1', 'Телефон-2', 'Телефон-3', 'Почта']
parsed_inner_numbers = ['Внутренний телефон-1', 'Внутренний телефон-2', 'Внутренний телефон-3']

def convert_data(row, parsed_numbers):
    phone = row['Телефон-1']
    if phone != '':
        phone = phone.split(' ~ ')
        for part in range(len(phone)):
            phone[part] = phone[part].replace(' ', '')
            phone[part] = phone[part].replace(u'\xa0', '')
            if phone[part].startswith('('):
                row[parsed_numbers[part]] = '+7' + (''.join(''.join(''.join(phone[part].split('(')).split(')')).split('-'))[:10])
                
            elif '@' in phone[part]:
                row['Почта'] = phone[part]
            else:
                pass
    inner_phone = row['Внутренний телефон-1'].split(' ~ ')
    for j in range(len(inner_phone)):
        row[parsed_inner_numbers[j]] = inner_phone[j]
    row['Адрес'] = row['Адрес'].replace('(', '')
    row['Адрес'] = row['Адрес'].replace(')', ',')
    if row['Телефон-1'] in row['Почта'] or row['Почта'] in row['Телефон-1']:
        row['Телефон-1'] = ''
    return row

reference_df = reference_df.apply(convert_data, args=[parsed_numbers], axis=1)

reference_df

import openpyxl
reference_df.to_excel('telephone_reference.xlsx')



